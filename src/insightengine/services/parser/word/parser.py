from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, TYPE_CHECKING

from docx import Document

if TYPE_CHECKING:  # pragma: no cover - 仅用于类型检查
    from docx.document import Document as DocxDocument

from ..base import Parser, ParserError
from ..registry import register_parser
from ..text_utils import build_items_from_chunks, normalize_whitespace
from ..types import ParseItem, ParseResult


_DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
_DOC_CONTENT_TYPE = "application/msword"


class DocParser(Parser):
    """解析 Microsoft Word `.doc` 旧格式文件。"""

    name = "doc"

    def parse(self, source: Any, **opts) -> ParseResult:
        data, source_path = _load_doc_bytes(source)
        chunks = _extract_legacy_doc_chunks(data)
        items = build_items_from_chunks(
            chunks,
            metadata_factory=lambda idx, _: {"paragraph": idx},
        )

        metadata_payload = {
            "parser": self.name,
            "content_type": _DOC_CONTENT_TYPE,
            "legacy_format": True,
            "chunk_count": len(items),
        }
        extra_metadata = opts.get("metadata", {})
        metadata_payload.update(extra_metadata)

        return ParseResult(
            source=source_path,
            items=items,
            metadata=metadata_payload,
        )


class DocxParser(Parser):
    """解析 Microsoft Word `.docx` 文件，输出结构化结果。"""

    name = "docx"

    def parse(self, source: Any, **opts) -> ParseResult:
        # 根据输入类型加载文档对象，同时返回可选的源路径
        document, source_path = _load_document(source)

        items = []
        position = 1

        for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
            text = normalize_whitespace(paragraph.text)
            if not text:
                continue
            items.append(
                ParseItem(
                    id=f"para-{paragraph_index}",
                    text=text,
                    length=len(text),
                    position=position,
                    metadata={"paragraph": paragraph_index},
                )
            )
            position += 1

        for table_index, table in enumerate(document.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                cell_text = normalize_whitespace(
                    " ".join(cell.text for cell in row.cells)
                )
                if not cell_text:
                    continue
                items.append(
                    ParseItem(
                        id=f"table-{table_index}-row-{row_index}",
                        text=cell_text,
                        length=len(cell_text),
                        position=position,
                        metadata={"table": table_index, "row": row_index},
                    )
                )
                position += 1

        metadata_payload = {
            "parser": self.name,
            "content_type": _DOCX_CONTENT_TYPE,
        }
        extra_metadata = opts.get("metadata", {})
        metadata_payload.update(extra_metadata)

        return ParseResult(
            source=source_path,
            items=items,
            metadata=metadata_payload,
        )


def _load_doc_bytes(source: Any) -> tuple[bytes, str | None]:
    """根据 source 类型读取 `.doc` 二进制数据。"""

    if isinstance(source, Path):
        return source.read_bytes(), str(source)

    if isinstance(source, str):
        path = Path(source)
        if path.exists():
            return path.read_bytes(), str(path)
        raise ParserError("string source for doc must be an existing file path")

    if isinstance(source, (bytes, bytearray)):
        return bytes(source), None

    if hasattr(source, "read"):
        data = source.read()
        if isinstance(data, str):
            data = data.encode("utf-8")
        if not isinstance(data, (bytes, bytearray)):
            raise ParserError("file-like object for doc must return bytes or str")
        return bytes(data), getattr(source, "name", None)

    raise ParserError(f"unsupported DOC source type: {type(source)!r}")


def _extract_legacy_doc_chunks(data: bytes) -> list[str]:
    """从 `.doc` 二进制数据中粗粒度提取可读文本。"""

    # `.doc` 文件通常以 UTF-16 LE 存储正文文本，这里先尝试直接解码
    unicode_text = data.decode("utf-16-le", errors="ignore")
    unicode_chunks = _split_control_sequences(unicode_text)

    ascii_text = data.decode("latin-1", errors="ignore")
    ascii_chunks = _split_control_sequences(ascii_text)

    if not unicode_chunks and ascii_chunks:
        return ascii_chunks
    if not ascii_chunks:
        return unicode_chunks

    if _chunk_readability_score(ascii_chunks) > _chunk_readability_score(
        unicode_chunks
    ):
        return ascii_chunks
    return unicode_chunks


def _split_control_sequences(text: str) -> list[str]:
    parts = [part.strip() for part in re.split(r"[\x00-\x09\x0b-\x1f]+", text)]
    return [part for part in parts if part]


def _chunk_readability_score(chunks: list[str]) -> float:
    text = "".join(chunks)
    base = sum(1 for ch in text if ch.isprintable() and not ch.isspace())
    penalty = len(chunks) * 0.5
    return base - penalty


def _load_document(source: Any) -> tuple["DocxDocument", str | None]:
    """根据 source 类型加载 docx 文档，并尽可能返回原始路径。"""

    if isinstance(source, Path):
        return Document(str(source)), str(source)

    if isinstance(source, str):
        path = Path(source)
        if path.exists():
            return Document(str(path)), str(path)
        raise ParserError("string source for docx must be an existing file path")

    if isinstance(source, (bytes, bytearray)):
        return Document(io.BytesIO(bytes(source))), None

    if hasattr(source, "read"):
        # 类文件对象可能返回 str 或 bytes，这里统一处理并保留名称信息
        data = source.read()
        if isinstance(data, str):
            data = data.encode("utf-8")
        if not isinstance(data, (bytes, bytearray)):
            raise ParserError("file-like object for docx must return bytes or str")
        return Document(io.BytesIO(bytes(data))), getattr(source, "name", None)

    raise ParserError(f"unsupported DOCX source type: {type(source)!r}")


register_parser(DocParser.name, DocParser)
register_parser(DocxParser.name, DocxParser)

__all__ = ["DocParser", "DocxParser"]
